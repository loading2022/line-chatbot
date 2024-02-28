from flask import Flask, request, abort
from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError
from linebot.models import *
import os
import requests
import json
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.callbacks import get_openai_callback
from langchain.chat_models import ChatOpenAI
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
#from doc2docx import convert

app = Flask(__name__)

openai_api_key=os.getenv('OPENAI_API_KEY',None)
channel_secret = os.getenv('LINE_CHANNEL_SECRET', None)
channel_access_token = os.getenv('LINE_CHANNEL_ACCESS_TOKEN', None)

if channel_secret is None:
    print('Specify LINE_CHANNEL_SECRET as environment variable.')
    sys.exit(1)
if channel_access_token is None:
    print('Specify LINE_CHANNEL_ACCESS_TOKEN as environment variable.')
    sys.exit(1)
line_bot_api = LineBotApi(channel_access_token)
handler = WebhookHandler(channel_secret)

#file_contents = {}
text=""
def get_text_from_pdf(pdf_path):
    text = ""
    pdf_reader = PdfReader(BytesIO(pdf_path))
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def get_text_from_docx(docx_path):
    #doc = Document(docx_path)
    doc = Document(BytesIO(docx_path))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

@app.route("/callback", methods=['POST'])
def callback():
    body = request.get_data(as_text=True)
    json_data = json.loads(body)
    print(json_data)               # 印出 json_data

    
    signature = request.headers['X-Line-Signature']
    handler.handle(body, signature)
    return 'OK'

@handler.add(MessageEvent, message=FileMessage)
def handle_file_message(event):
    global text
    #user_id = event['events'][0]['source']['userId']
    file_id=event.message.id
    #file_contents[user_id] = ""
    url = f"https://api-data.line.me/v2/bot/message/{file_id}/content"
    headers = {
        'Authorization': f'Bearer {channel_access_token}'
    }
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        file_name = event.message.file_name
        file_content = response.content
        if file_name.endswith('.pdf'):
            #file_contents[user_id] += get_text_from_pdf(file_content)
            text+= get_text_from_pdf(file_content)
        elif file_name.endswith('.docx'):
            #file_contents[user_id] += get_text_from_docx(file_content)
            text+= get_text_from_docx(file_content)
        elif file_name.endswith('.txt'):
            file_content = response.content.decode('utf-8')
            #file_contents[user_id] += file_content
            text+= file_content
        line_bot_api.reply_message(event.reply_token, TextSendMessage(text="文件上傳成功!\n可以開始問相關問題"))   
    else:
        line_bot_api.reply_message(event.reply_token, TextSendMessage(text="讀取文件失敗"))

@handler.add(MessageEvent, message=TextMessage)
def handle_text_message(event):
    global text
    #user_id = event['events'][0]['source']['userId']
    user_message = event.message.text
    print(user_message)
    if user_message=="開啟新對話":
        #file_contents[user_id] = ""
        text=""
        line_bot_api.reply_message(event.reply_token, TextSendMessage(text='已開啟新對話\n可重新上傳檔案，目前支援 pdf、docx 和 txt 檔'))
    else:
        #if file_contents[user_id] =="":
        #    line_bot_api.reply_message(event.reply_token, TextSendMessage(text='請上傳檔案，目前支援 pdf、docx 和 txt 檔'))
        if text=="":
            line_bot_api.reply_message(event.reply_token, TextSendMessage(text='請上傳檔案，目前支援 pdf、docx 和 txt 檔'))
        else:
            text_splitter = CharacterTextSplitter(
                    separator="\n",
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len
            )
            chunks = text_splitter.split_text(file_contents[user_id])
        
            embeddings = OpenAIEmbeddings()
            knowledge_base = FAISS.from_texts(chunks, embeddings)
        
            docs = knowledge_base.similarity_search(user_message)
            llm = ChatOpenAI(
                model_name="gpt-4-1106-preview",
                temperature=0.4
            )
        
            chain = load_qa_chain(llm, chain_type="stuff")
        
            with get_openai_callback() as cb:
                response = chain.run(input_documents=docs, question=user_message)
        
            response = {'response': response}
        
            line_bot_api.reply_message(event.reply_token, TextSendMessage(text=response['response']))

if __name__ == "__main__":
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
